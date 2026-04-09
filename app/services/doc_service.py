from io import BytesIO
from docx import Document
from docx.shared import Pt
from app.services.pauta_service import load_pauta


def build_pauta_docx():
    pauta = load_pauta()

    doc = Document()

    # Título
    title = doc.add_paragraph()
    title_run = title.add_run("PAUTA GRAN DÍA")
    title_run.bold = True
    title_run.font.size = Pt(16)

    doc.add_paragraph("")  # línea en blanco

    for row in pauta:
        numero = str(row.get("numero", "")).strip()
        nombre = str(row.get("nombre", "")).strip().upper()
        texto = str(row.get("texto", "")).strip()
        tipo = str(row.get("tipo", "FOTO")).strip().upper()
        if tipo not in {"FOTO", "CUBRIR", "TOTAL"}:
            tipo = "FOTO"

        linea = f"{numero} - {nombre}"
        if texto:
            linea += f" - {texto}"
        linea += f" - {tipo}"

        p = doc.add_paragraph()
        run = p.add_run(linea)
        run.font.size = Pt(12)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
